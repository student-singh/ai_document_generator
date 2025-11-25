# utils/docx_generator.py
from docx import Document
import os

def create_docx(project):
    doc = Document()
    doc.add_heading(project["topic"], 0)
    
    for i, title in enumerate(project["structure"]):
        doc.add_heading(title, level=1)
        text = project["content"].get(str(i), {}).get("text", "")
        if text.strip():
            doc.add_paragraph(text)
        else:
            doc.add_paragraph("Content not generated yet.")
    
    os.makedirs("uploads", exist_ok=True)
    filepath = f"uploads/{project['_id']}.docx"
    doc.save(filepath)
    return filepath